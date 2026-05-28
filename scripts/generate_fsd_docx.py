"""
Generate standalone Farmer Segmentation FSD Word documents (BAT-style).
Output: docs/fsd/*.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT_DIR = Path(__file__).resolve().parents[1] / "docs" / "fsd"
IMAGES_DIR = Path(__file__).resolve().parents[1] / "docs" / "images"

PROJECT = "MVP 2 Farming Engagement Hub"
EPIC = "Farmer Segmentation"
PORTAL = "Back Office PA (Web)"
PROTOTYPE = "BAT - Agri 360 — Figma (segmentation flows)"
AUTHOR_NAME = "Matheus Kist"
OPERATOR_ROLE = "Backoffice PA User"
WIREFRAME_WIDTH = Inches(6.25)


def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _wireframe(doc: Document, filename: str, caption: str | None = None) -> None:
    path = IMAGES_DIR / filename
    if not path.is_file():
        _para(doc, f"[Wireframe not found: {filename}]", bold=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=WIREFRAME_WIDTH)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)
    doc.add_paragraph()


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def _cover_block(doc: Document, meta: dict) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Functional Specification Document")
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph()

    fields = [
        ("PROJECT", PROJECT),
        ("EPIC", EPIC),
        ("Functionality", meta["functionality"]),
        ("User Story ID", meta["story_id"]),
        ("Functional Specification For", "(X) Web Application    ( ) Mobile Application"),
        ("Portal", PORTAL),
        ("Prototype", PROTOTYPE),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)
    doc.add_page_break()


def _standard_sections(doc: Document, meta: dict, sections: dict) -> None:
    _heading(doc, "Contents", 1)
    for i, title in enumerate(
        [
            "1. Revision History",
            "2. Sign-off",
            "3. Process Specification",
            "4. User Interface & User Experience",
            "5. Policy & Regulatory Requirements",
            "6. Technical Requirements",
            "7. Positive & Negative Testing",
        ],
        start=1,
    ):
        doc.add_paragraph(title, style="List Number")
    doc.add_page_break()

    _heading(doc, "1. Revision History", 1)
    _table(
        doc,
        ["Date", "Name", "Change", "Version"],
        [["26/05/2026", AUTHOR_NAME, "Initial document draft", "1.0"]],
    )

    _heading(doc, "2. Sign-off", 1)
    _para(
        doc,
        f"All screens described in this document are operated by the {OPERATOR_ROLE} "
        "within the Back Office PA portal.",
    )
    _table(
        doc,
        ["Date", "Name", "Role"],
        [
            ["", "", ""],
            ["", "", ""],
            ["", "", ""],
        ],
    )

    for key, (title, level) in [
        ("process", ("3. Process Specification", 1)),
        ("ui", ("4. User Interface & User Experience", 1)),
        ("policy", ("5. Policy & Regulatory Requirements", 1)),
        ("technical", ("6. Technical Requirements", 1)),
        ("testing", ("7. Positive & Negative Testing", 1)),
    ]:
        _heading(doc, title, level)
        body = sections[key]
        if isinstance(body, list):
            for block in body:
                if block["type"] == "heading":
                    _heading(doc, block["text"], block.get("level", 2))
                elif block["type"] == "para":
                    _para(doc, block["text"], block.get("bold", False))
                elif block["type"] == "bullets":
                    _bullets(doc, block["items"])
                elif block["type"] == "numbered":
                    _numbered(doc, block["items"])
                elif block["type"] == "table":
                    _table(doc, block["headers"], block["rows"])
                elif block["type"] == "image":
                    _wireframe(doc, block["file"], block.get("caption"))
        else:
            _para(doc, body)


def _build_doc(meta: dict, sections: dict) -> Document:
    doc = Document()
    _set_normal_style(doc)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    _cover_block(doc, meta)
    _standard_sections(doc, meta, sections)
    return doc


# --- Content definitions (standalone, no API paths) ---

def fsd_001() -> tuple[dict, dict]:
    meta = {
        "story_id": "PA-SEG-001",
        "functionality": "Farmer Contract KPI — Processing and Registration",
        "filename": "PA-SEG-001 - Farmer Contract KPI - Processing and Registration.docx",
    }
    sections = {
        "process": [
            {
                "type": "heading",
                "text": "3.1 Business Process",
                "level": 2,
            },
            {
                "type": "para",
                "text": "In this section you can find the detailed business process description and the corresponding workflows.",
            },
            {
                "type": "bullets",
                "items": [
                    "What? Maintain Farmer Contract KPI facts per farmer, crop season, and tobacco culture type. "
                    "These facts feed segmentation simulations (Loyalty, Quality, Financial, Yield, Scale, and ESG percentage components).",
                    f"Who? {OPERATOR_ROLE}.",
                    "When? Before running segmentation simulations for a crop season; when contract or agronomic facts must be reflected in segmentation.",
                    "Where? Agri360 Web — KPI Data area, Farmer Contract KPI screen.",
                    "Why? Segmentation scoring requires a single, auditable store of contract-level KPI values.",
                    "How? Users filter by crop season, review a grid, edit rows manually, and/or import a bulk file. "
                    "In a later phase, scheduled jobs will calculate KPIs from integrated sources and freeze values in the same store.",
                ],
            },
            {
                "type": "heading",
                "text": "3.2 Phased behaviour",
                "level": 2,
            },
            {
                "type": "table",
                "headers": ["Phase", "Behaviour"],
                "rows": [
                    [
                        "Phase 1 (MVP)",
                        "The screen is the authorised channel for data load: manual grid maintenance and file import. "
                        "No automatic calculation from contracting or commercialization modules yet.",
                    ],
                    [
                        "Phase 2 (target Agri360)",
                        "Background routines compute contract KPIs from integrated sources. Results are frozen in the "
                        "Farmer Contract KPI store with calculation metadata. The UI shows calculated rows as read-only "
                        "unless the user has override permission; overrides are audited.",
                    ],
                ],
            },
            {
                "type": "heading",
                "text": "3.3 Data requirements",
                "level": 2,
            },
            {
                "type": "para",
                "text": "One logical record per Farmer + Crop Season + Culture Type (composite key).",
            },
            {
                "type": "table",
                "headers": ["Field", "Required", "Description"],
                "rows": [
                    ["Farmer code", "Yes", "Business key referencing Farmer master."],
                    ["Crop season", "Yes", "Crop season identifier (e.g. year 2025)."],
                    ["Culture type code", "No", "Tobacco culture type; default FCV when omitted."],
                    ["Delivered percentage", "Yes", "Delivered % of contracted volume."],
                    ["Delivered amount (kg)", "Yes", "Delivered volume in kilograms."],
                    ["Contracted amount (kg)", "Yes", "Contracted volume in kilograms."],
                    ["IQS", "Yes", "Integrated Quality Score."],
                    ["Had NTRM", "Yes", "Whether farmer had NTRM event."],
                    ["Had quality mixture", "Yes", "Whether farmer had quality mixture issue."],
                    ["Self-funding percentage", "Yes", "Self-funding %."],
                    ["Have debt", "Yes", "Whether farmer has debt."],
                    ["Yield", "Yes", "Yield metric for yield KPI scoring."],
                    ["Scale (ha)", "Yes", "Planted scale in hectares; used to resolve culture type in simulations."],
                    ["Reforestation percentage", "Yes", "ESG reforestation %."],
                    ["Native forest percentage", "Yes", "ESG native forest %."],
                    ["Non-exclusive flag", "Derived", "Snapshot of farmer non-exclusive status at save/import time."],
                ],
            },
            {
                "type": "para",
                "text": "Import rule: If delivered amount is omitted, the system may derive it from delivered percentage and contracted amount.",
            },
        ],
        "ui": [
            {
                "type": "para",
                "text": "In this section you can find details on screens, navigation, CRUD behaviour, validations, and error handling.",
            },
            {
                "type": "heading",
                "text": "4.1 General rules",
                "level": 2,
            },
            {
                "type": "bullets",
                "items": [
                    "A global crop season selector filters the grid and sets the default season for new rows and import.",
                    "Rows are maintained per culture type.",
                    f"Access is limited to the {OPERATOR_ROLE} (view, edit, import; Phase 2 override where applicable).",
                    "All manual changes and imports must be audited (user, timestamp, before/after values).",
                ],
            },
            {
                "type": "heading",
                "text": "4.2 SCREEN 01: KPI Data — Farmer Contract KPI (Web)",
                "level": 2,
            },
            {
                "type": "para",
                "text": "Screen purpose: List, manually maintain, and bulk-import farmer contract KPI rows for the selected crop season.",
            },
            {
                "type": "image",
                "file": "PA-SEG-001 - Wireframe.png",
                "caption": "Figure 1 — Wireframe: Farmer Contract KPI screen (tabs for contract, technologies, and irregularities on the same KPI Data area).",
            },
            {
                "type": "heading",
                "text": "4.2.1 Layout",
                "level": 3,
            },
            {
                "type": "bullets",
                "items": [
                    "Header: page title, crop season selector, contextual help.",
                    "Toolbar: Import file, Add row, optional Export.",
                    "Grid: all data fields listed in section 3.3 for the active season.",
                    "Status area: import summary (inserted, updated, errors) with drill-down to row-level messages.",
                ],
            },
            {
                "type": "heading",
                "text": "4.2.2 Manual maintenance",
                "level": 3,
            },
            {
                "type": "bullets",
                "items": [
                    "Add row: validate farmer exists; enforce unique Farmer + Season + Culture.",
                    "Edit: inline or form edit with field-level validation.",
                    "Delete: confirmation required; removes KPI record for that key.",
                ],
            },
            {
                "type": "heading",
                "text": "4.2.3 File import",
                "level": 3,
            },
            {
                "type": "bullets",
                "items": [
                    "Supported format: CSV (UTF-8) with header row.",
                    "Column headers are matched case-insensitively; spaces and underscores ignored.",
                    "Upsert key: Farmer code + Crop season + Culture type.",
                    "Partial success: valid rows commit; invalid rows reported with row number and message.",
                    "Required columns: Farmer code, Crop season, Delivered percentage, Contracted amount (kg), IQS, "
                    "Had NTRM, Had quality mixture, Self-funding %, Have debt, Yield, Scale, Reforestation %, Native forest %.",
                    "Optional columns: Culture type, Delivered amount (kg).",
                ],
            },
            {
                "type": "heading",
                "text": "4.2.4 Phase 2 UI",
                "level": 3,
            },
            {
                "type": "bullets",
                "items": [
                    "Source column: Manual / Calculated / Override.",
                    "Calculated fields read-only unless override role is granted.",
                    "Display last calculation timestamp and batch reference.",
                ],
            },
            {
                "type": "heading",
                "text": "4.3 Dependencies on predecessor user stories",
                "level": 2,
            },
            {
                "type": "bullets",
                "items": [
                    "Farmer master data and crop season catalogue.",
                    "Culture type catalogue.",
                    "Segmentation configuration and simulation user stories consume this data once loaded.",
                ],
            },
        ],
        "policy": [
            {
                "type": "para",
                "text": "Segregation of duties: users who import or edit KPI data should not be sole approvers of official segmentation.",
            },
            {
                "type": "para",
                "text": "Audit trail required for all KPI changes and imports. Farmer personal data in grids must follow data privacy policies.",
            },
        ],
        "technical": [
            {
                "type": "heading",
                "text": "6.1 Automations",
                "level": 2,
            },
            {
                "type": "bullets",
                "items": [
                    "Phase 1: none (manual and import only).",
                    "Phase 2: scheduled job to compute and freeze contract KPIs from Agri360 source data per season; notify operations on failure.",
                ],
            },
        ],
        "testing": [
            {
                "type": "heading",
                "text": "7.1 Test scenarios and acceptance criteria (Web)",
                "level": 2,
            },
            {
                "type": "table",
                "headers": ["ID", "Scenario", "Expected result"],
                "rows": [
                    ["T-01", "Import valid file for a crop season", "Rows upserted; summary shows counts."],
                    ["T-02", "Import unknown farmer code", "Row error; other valid rows succeed."],
                    ["T-03", "Manual edit delivered %", "Saved value used in subsequent simulation."],
                    ["T-04", "Delete row", "Row removed from KPI store."],
                    ["T-05", "Phase 2 calculated row", "Fields read-only without override role."],
                ],
            },
        ],
    }
    return meta, sections


def fsd_002() -> tuple[dict, dict]:
    meta = {
        "story_id": "PA-SEG-002",
        "functionality": "Farmer Technologies KPI — Processing and Registration",
        "filename": "PA-SEG-002 - Farmer Technologies KPI - Processing and Registration.docx",
    }
    sections = {
        "process": [
            {
                "type": "bullets",
                "items": [
                    "What? Register technology adoptions per farmer, crop season, and culture type for Technology KPI scoring.",
                    f"Who? {OPERATOR_ROLE}.",
                    "When? Before simulations where technology adoption affects scores.",
                    "Where? KPI Data — Technologies screen (Web).",
                    "Why? Technology facts are separate from contract KPI rows (multiple technologies per farmer per season).",
                    "How? Grid, manual maintenance, and file import (Phase 1); Agri360-fed data (Phase 2).",
                ],
            },
            {
                "type": "heading",
                "text": "3.2 Phased behaviour",
                "level": 2,
            },
            {
                "type": "table",
                "headers": ["Phase", "Behaviour"],
                "rows": [
                    [
                        "Phase 1 (MVP)",
                        "Manual grid maintenance and file import only.",
                    ],
                    [
                        "Phase 2 (target Agri360)",
                        "Technology adoption data collected by Agri360 (e.g. field notebook, BAT tech package) "
                        "is written to this KPI store through scheduled routines. Rows are frozen with calculation "
                        "metadata; the UI is read-only except for authorised overrides.",
                    ],
                ],
            },
            {
                "type": "heading",
                "text": "3.3 Data requirements",
                "level": 2,
            },
            {
                "type": "para",
                "text": "One record per Farmer + Crop Season + Culture Type + Technology.",
            },
            {
                "type": "table",
                "headers": ["Field", "Required", "Description"],
                "rows": [
                    ["Farmer code", "Yes", "References Farmer master."],
                    ["Crop season", "Yes", "Season identifier."],
                    ["Culture type code", "No", "Default FCV."],
                    ["Technology", "Yes", "Reference to Technology catalogue."],
                ],
            },
        ],
        "ui": [
            {
                "type": "heading",
                "text": "4.1 SCREEN 01: KPI Data — Technologies (Web)",
                "level": 2,
            },
            {
                "type": "image",
                "file": "PA-SEG-002 - Wireframe.png",
                "caption": "Figure 1 — Wireframe: Technologies KPI screen.",
            },
            {
                "type": "bullets",
                "items": [
                    "Grid columns: farmer code, farmer name, crop season, culture type, technology name.",
                    "Toolbar: Import, Add row, Export (optional).",
                    "Add: validate catalogue technology; prevent duplicate natural key.",
                    "Delete: removes adoption for that season.",
                    "Import columns: Farmer code, Crop season, Culture type (optional), Technology identifier.",
                ],
            },
            {
                "type": "para",
                "text": "Help text: Segmentation sums configured technology scores for technologies present in the farmer's latest active season within the simulation Technologies scope (see PA-SEG-005).",
            },
        ],
        "policy": [
            {"type": "para", "text": "Audit all changes. Segregation of duties vs segmentation approval."},
        ],
        "technical": [
            {
                "type": "heading",
                "text": "6.1 Automations",
                "level": 2,
            },
            {
                "type": "bullets",
                "items": [
                    "Phase 1: none (manual and import only).",
                    "Phase 2: scheduled sync from Agri360 technology adoption sources.",
                ],
            },
        ],
        "testing": [
            {
                "type": "table",
                "headers": ["ID", "Scenario", "Expected result"],
                "rows": [
                    ["T-01", "Two technologies same farmer/season", "Both contribute to Technology score if configured."],
                    ["T-02", "Invalid technology reference", "Import row error."],
                    ["T-03", "Delete adoption", "Technology no longer scores in simulation."],
                ],
            },
        ],
    }
    return meta, sections


def fsd_003() -> tuple[dict, dict]:
    meta = {
        "story_id": "PA-SEG-003",
        "functionality": "Farmer Irregularities KPI — Processing and Registration",
        "filename": "PA-SEG-003 - Farmer Irregularities KPI - Processing and Registration.docx",
    }
    sections = {
        "process": [
            {
                "type": "bullets",
                "items": [
                    "What? Record ESG/compliance irregularity events per farmer, crop season, and culture type.",
                    f"Who? {OPERATOR_ROLE}.",
                    "When? When irregularities are known and before ESG scoring simulations.",
                    "Where? KPI Data — ESG Irregularities screen (Web); user story name: Farmer Irregularities KPI.",
                    "Why? Irregularities are multi-valued and separate from contract KPI percentages.",
                    "How? Grid, manual maintenance, import (Phase 1); Agri360-fed data (Phase 2).",
                ],
            },
            {
                "type": "heading",
                "text": "3.2 Phased behaviour",
                "level": 2,
            },
            {
                "type": "table",
                "headers": ["Phase", "Behaviour"],
                "rows": [
                    [
                        "Phase 1 (MVP)",
                        "Manual grid maintenance and file import only.",
                    ],
                    [
                        "Phase 2 (target Agri360)",
                        "Irregularity events captured by Agri360 compliance and ESG processes are written to this "
                        "KPI store through scheduled routines. Rows are frozen with calculation metadata; the UI is "
                        "read-only except for authorised overrides.",
                    ],
                ],
            },
            {
                "type": "heading",
                "text": "3.3 Scoring link",
                "level": 2,
            },
            {
                "type": "para",
                "text": "Segmentation configuration assigns a score per irregularity type. At simulation, irregularities in the latest active ESG scope season add to the ESG component together with reforestation and native forest percentages from Farmer Contract KPI.",
            },
            {
                "type": "heading",
                "text": "3.4 Data requirements",
                "level": 2,
            },
            {
                "type": "table",
                "headers": ["Field", "Required", "Description"],
                "rows": [
                    ["Farmer code", "Yes", "Farmer business key."],
                    ["Crop season", "Yes", "Crop season."],
                    ["Culture type", "No", "Default FCV."],
                    ["Irregularity type", "Yes", "Catalogue reference."],
                ],
            },
        ],
        "ui": [
            {
                "type": "heading",
                "text": "4.1 SCREEN 01: KPI Data — ESG Irregularities (Web)",
                "level": 2,
            },
            {
                "type": "image",
                "file": "PA-SEG-003 - Wireframe.png",
                "caption": "Figure 1 — Wireframe: ESG irregularities (Farmer Irregularities KPI) screen.",
            },
            {
                "type": "bullets",
                "items": [
                    "Same interaction pattern as Technologies: grid, add, delete, import.",
                    "Import columns: Farmer code, Crop season, Culture type (optional), Irregularity type.",
                    "Tooltip: reforestation/native forest % maintained on Farmer Contract KPI; events maintained here.",
                ],
            },
        ],
        "policy": [
            {
                "type": "para",
                "text": "ESG data may be sensitive — restrict by role. Align with environmental compliance policies (Legal TBD).",
            },
        ],
        "technical": [
            {
                "type": "heading",
                "text": "6.1 Automations",
                "level": 2,
            },
            {
                "type": "bullets",
                "items": [
                    "Phase 1: none (manual and import only).",
                    "Phase 2: scheduled sync from Agri360 compliance / ESG case closure.",
                ],
            },
        ],
        "testing": [
            {
                "type": "table",
                "headers": ["ID", "Scenario", "Expected result"],
                "rows": [
                    ["T-01", "Import irregularity", "ESG simulation includes configured score when type matches."],
                    ["T-02", "Multiple irregularities", "Scores aggregate per configuration rules."],
                    ["T-03", "No irregularities", "Only reforestation/native forest portion applies."],
                ],
            },
        ],
    }
    return meta, sections


def fsd_004() -> tuple[dict, dict]:
    meta = {
        "story_id": "PA-SEG-004",
        "functionality": "Segmentation Configuration",
        "filename": "PA-SEG-004 - Segmentation Configuration.docx",
    }
    sections = {
        "process": [
            {
                "type": "bullets",
                "items": [
                    "What? Define segmentation configurations: named scoring rules and segment tiers for farmer simulations.",
                    f"Who? {OPERATOR_ROLE}.",
                    "When? Before first simulation of a policy cycle; when tiers or KPI weights change.",
                    "Where? Segmentation Configurations list and editor (Web).",
                    "Why? Reproducible, comparable scoring policies.",
                    "How? Create/edit configuration, header segments, per-culture-type KPI blocks and thresholds; save blocked until validations pass.",
                ],
            },
            {
                "type": "para",
                "text": "Configurations are crop-season-agnostic. Crop seasons are selected when a simulation is run (PA-SEG-005).",
            },
            {
                "type": "heading",
                "text": "3.2 Logical structure",
                "level": 2,
            },
            {
                "type": "bullets",
                "items": [
                    "Configuration name.",
                    "Header segments: name, only-exclusive-farmer flag, bank deposit discount, tobacco discount.",
                    "Per culture type (FCV, BLY, AM, …): maximum score, segment score thresholds, seven KPI blocks (Loyalty, Quality, Financial, Technology, ESG, Yield, Scale).",
                    "Each KPI block: max score, relevance % (read-only: KPI max ÷ culture max), and rule rows.",
                ],
            },
        ],
        "ui": [
            {
                "type": "heading",
                "text": "4.1 SCREEN 01: Configurations — List",
                "level": 2,
            },
            {
                "type": "bullets",
                "items": [
                    "Columns: name, culture types included.",
                    "Actions: Create, Edit, Duplicate (copies full rule set under new name).",
                ],
            },
            {
                "type": "image",
                "file": "PA-SEG-004 - List.png",
                "caption": "Figure 1 — Wireframe: Segmentation configurations list.",
            },
            {
                "type": "heading",
                "text": "4.2 SCREEN 02: Configuration editor",
                "level": 2,
            },
            {
                "type": "heading",
                "text": "4.2.1 Header — Configuration name and segments",
                "level": 3,
            },
            {
                "type": "image",
                "file": "PA-SEG-004 - Header.png",
                "caption": "Figure 2 — Wireframe: configuration header (name and segment definitions).",
            },
            {
                "type": "bullets",
                "items": [
                    "At least one segment required.",
                    "Range min (per culture type tab): minimum total simulation score to qualify; may be negative.",
                    "Empty range min: catch-all tier after scored thresholds.",
                    "Only exclusive farmer: non-exclusive farmers cannot receive that segment in simulation.",
                ],
            },
            {
                "type": "heading",
                "text": "4.2.2 Culture type — maximum score and segment thresholds",
                "level": 3,
            },
            {
                "type": "image",
                "file": "PA-SEG-004 - CultureType Config.png",
                "caption": "Figure 3 — Wireframe: culture type block (maximum score and segment thresholds).",
            },
            {
                "type": "heading",
                "text": "4.2.3 KPI blocks per culture type",
                "level": 3,
            },
            {
                "type": "para",
                "text": "Each culture type tab contains seven KPI configuration areas. Wireframes below illustrate each KPI editor.",
            },
            {
                "type": "heading",
                "text": "4.2.3.1 Loyalty",
                "level": 4,
            },
            {
                "type": "image",
                "file": "PA-SEG-004 - Loyalty.png",
                "caption": "Figure 4 — Wireframe: Loyalty KPI configuration.",
            },
            {
                "type": "heading",
                "text": "4.2.3.2 Quality",
                "level": 4,
            },
            {
                "type": "image",
                "file": "PA-SEG-004 - Quality.png",
                "caption": "Figure 5 — Wireframe: Quality KPI configuration.",
            },
            {
                "type": "heading",
                "text": "4.2.3.3 Financial",
                "level": 4,
            },
            {
                "type": "image",
                "file": "PA-SEG-004 - Financials.png",
                "caption": "Figure 6 — Wireframe: Financial KPI configuration.",
            },
            {
                "type": "heading",
                "text": "4.2.3.4 Technology",
                "level": 4,
            },
            {
                "type": "image",
                "file": "PA-SEG-004 - Technology.png",
                "caption": "Figure 7 — Wireframe: Technology KPI configuration.",
            },
            {
                "type": "heading",
                "text": "4.2.3.5 ESG",
                "level": 4,
            },
            {
                "type": "image",
                "file": "PA-SEG-004 - ESG.png",
                "caption": "Figure 8 — Wireframe: ESG KPI configuration.",
            },
            {
                "type": "heading",
                "text": "4.2.3.6 Yield",
                "level": 4,
            },
            {
                "type": "image",
                "file": "PA-SEG-004 - Yield.png",
                "caption": "Figure 9 — Wireframe: Yield KPI configuration.",
            },
            {
                "type": "heading",
                "text": "4.2.3.7 Scale",
                "level": 4,
            },
            {
                "type": "image",
                "file": "PA-SEG-004 - Scale.png",
                "caption": "Figure 10 — Wireframe: Scale KPI configuration.",
            },
            {
                "type": "table",
                "headers": ["KPI block", "Configuration content"],
                "rows": [
                    ["Loyalty", "Season-quantity rules (planting count + delivery window); historical volume on consolidated delivered %."],
                    ["Quality", "IQS ranges with look-back; scores for NTRM and quality mixture flags."],
                    ["Financial", "Self-funding % ranges; score for have-debt flag."],
                    ["Technology", "Score per technology from catalogue."],
                    ["ESG", "Reforestation/native forest points per % with caps; score per irregularity type."],
                    ["Yield", "Numeric yield ranges with look-back."],
                    ["Scale", "Numeric scale (ha) ranges with look-back."],
                ],
            },
            {
                "type": "heading",
                "text": "4.3 Business rules — Save validation",
                "level": 2,
            },
            {
                "type": "numbered",
                "items": [
                    "Sum of KPI max scores = culture type maximum score.",
                    "Each KPI max score must equal derived cap from its rules (see table below).",
                    "All seven KPI blocks required per culture type.",
                    "At least one header segment.",
                ],
            },
            {
                "type": "table",
                "headers": ["KPI", "Derived max score"],
                "rows": [
                    ["Loyalty", "Max positive season-quantity score + max positive historical volume score."],
                    ["Quality", "Maximum IQS range score (0 if no ranges)."],
                    ["Financial", "Maximum self-funding range score."],
                    ["Technology", "Sum of positive technology scores."],
                    ["ESG", "Reforestation cap + native forest cap + sum of positive irregularity scores."],
                    ["Yield", "Maximum yield range score."],
                    ["Scale", "Maximum scale range score."],
                ],
            },
            {
                "type": "heading",
                "text": "4.4 Loyalty precedence (simulation)",
                "level": 2,
            },
            {
                "type": "para",
                "text": "When multiple season-quantity rules match, the rule with the highest planting crop season amount wins. Consolidated delivered % uses sum of delivered kg ÷ sum of contracted kg across loyalty scope seasons — not an average of season percentages.",
            },
        ],
        "policy": [
            {
                "type": "para",
                "text": f"Restrict create/edit to {OPERATOR_ROLE}. Duplicate preserves historical simulations linked to original configuration id. Discount fields may affect fiscal reporting — confirm with Finance.",
            },
        ],
        "technical": [
            {
                "type": "bullets",
                "items": [
                    "Culture type, technology, and irregularity catalogues for pickers.",
                    "Full graph save under 5 seconds for typical three culture types.",
                    "Optimistic concurrency on update (target).",
                ],
            },
        ],
        "testing": [
            {
                "type": "table",
                "headers": ["ID", "Scenario", "Expected result"],
                "rows": [
                    ["T-01", "KPI sum ≠ maximum score", "Save rejected with clear message."],
                    ["T-02", "Technology max ≠ sum of tech scores", "Save rejected."],
                    ["T-03", "Duplicate configuration", "New configuration id; rules copied."],
                    ["T-04", "Valid 100-point configuration", "Save succeeds."],
                ],
            },
        ],
    }
    return meta, sections


def fsd_005() -> tuple[dict, dict]:
    meta = {
        "story_id": "PA-SEG-005",
        "functionality": "Segmentation Simulation and Approval",
        "filename": "PA-SEG-005 - Segmentation Simulation and Approval.docx",
    }
    sections = {
        "process": [
            {
                "type": "bullets",
                "items": [
                    "What? Run segmentation simulation; accept one result as official farmer segmentation for a target crop season — via Approval Workflow.",
                    f"Who? {OPERATOR_ROLE} runs simulations; designated approvers accept as official via Approval Workflow.",
                    "When? After KPI load and configuration publish.",
                    "Where? Simulations list, create form, simulation detail (Web).",
                    "How? Select configuration, target season, per-KPI scopes → run → review → submit approval → on approval, persist official segmentation.",
                ],
            },
            {
                "type": "heading",
                "text": "3.2 Simulation lifecycle",
                "level": 2,
            },
            {
                "type": "table",
                "headers": ["Status", "Meaning"],
                "rows": [
                    ["Simulation", "Completed run; not official."],
                    ["Official", "Accepted; drives official farmer segmentation for target season."],
                ],
            },
            {
                "type": "bullets",
                "items": [
                    "Multiple simulations per target season allowed (Simulation status).",
                    "At most one Official per target crop season.",
                    "Accepting one simulation demotes any other Official for that season.",
                    "Accept replaces all official segmentation rows for the season with simulation results (excluding new farmers).",
                ],
            },
            {
                "type": "heading",
                "text": "3.3 Simulation inputs",
                "level": 2,
            },
            {
                "type": "bullets",
                "items": [
                    "Segmentation configuration.",
                    "Target crop season.",
                    "Per-KPI scope: crop season list for each of Loyalty, Quality, Financial, ESG, Technologies, Yield, Scale.",
                    "Optional value aggregation (Average / Last active crop data) where applicable.",
                ],
            },
        ],
        "ui": [
            {
                "type": "heading",
                "text": "4.1 SCREEN 01: Simulations — List and create",
                "level": 2,
            },
            {
                "type": "bullets",
                "items": [
                    "Target crop season selector.",
                    "Configuration dropdown.",
                    "KPI scope section: multi-select seasons per KPI; aggregation radios per rules below.",
                    "Prominent new-farmer rule callout.",
                    "Run enabled when configuration and all seven KPI scopes are complete.",
                ],
            },
            {
                "type": "image",
                "file": "PA-SEG-005 - Run Simulation Sample.png",
                "caption": "Figure 1 — Wireframe: simulations list and new simulation (KPI scopes).",
            },
            {
                "type": "heading",
                "text": "4.2 SCREEN 02: Simulation detail",
                "level": 2,
            },
            {
                "type": "image",
                "file": "PA-SEG-005 - View Simulation Sample.png",
                "caption": "Figure 2 — Wireframe: simulation results and submit for approval.",
            },
            {
                "type": "bullets",
                "items": [
                    "Read-only KPI scope summary.",
                    "Segment distribution chart (excludes new farmers).",
                    "Farmer results: code, name, culture type, component scores, total, segment, new-farmer flag.",
                    "Export to spreadsheet.",
                    "Submit for approval (replaces direct Accept in production).",
                ],
            },
            {
                "type": "heading",
                "text": "4.3 Scoring rules — Overview",
                "level": 2,
            },
            {
                "type": "heading",
                "text": "4.3.1 New farmer gate",
                "level": 3,
            },
            {
                "type": "para",
                "text": "Prior season = target crop season year minus one. If the farmer has no Farmer Contract KPI for that prior season and culture type: all scores = 0, no segment, excluded from distribution chart. This gate uses target season only, not per-KPI scopes.",
            },
            {
                "type": "heading",
                "text": "4.3.2 Culture type selection",
                "level": 3,
            },
            {
                "type": "para",
                "text": "Use the culture type with the largest total Scale summed across Scale-scope seasons.",
            },
            {
                "type": "heading",
                "text": "4.3.3 Aggregation by KPI",
                "level": 3,
            },
            {
                "type": "table",
                "headers": ["KPI", "Aggregation", "Notes"],
                "rows": [
                    ["Loyalty", "N/A", "Multi-season rules on scope seasons."],
                    ["Quality", "Average or Last active", "IQS uses selection; NTRM and mixture always last active in scope."],
                    ["Financial", "Last active or Average", "Self-funding uses selection; have debt always last active."],
                    ["ESG", "N/A", "Latest active season in ESG scope."],
                    ["Technologies", "N/A", "Latest active season in Technologies scope."],
                    ["Yield", "Average or Last active", "Range match on aggregated yield."],
                    ["Scale", "Last active or Average", "Range match on aggregated scale (ha)."],
                ],
            },
            {
                "type": "para",
                "text": "Average = mean over scope seasons where farmer has data. Last active = value from highest season year in scope with data.",
            },
            {
                "type": "heading",
                "text": "4.3.4 Component rules (summary)",
                "level": 3,
            },
            {
                "type": "bullets",
                "items": [
                    "Loyalty: planting count, delivery window on last D seasons, consolidated delivered % for historical bands; highest planting rule wins.",
                    "Quality: IQS range match + NTRM/mixture add-ons.",
                    "Financial: self-funding range + debt add-on.",
                    "Technology: sum configured scores for adopted technologies in latest scope season.",
                    "ESG: reforestation and native forest points capped + irregularity scores.",
                    "Yield / Scale: range match on aggregated values.",
                ],
            },
            {
                "type": "heading",
                "text": "4.3.5 Segment assignment",
                "level": 3,
            },
            {
                "type": "numbered",
                "items": [
                    "Total score = sum of seven components.",
                    "Segments with range min evaluated highest threshold first.",
                    "Respect only-exclusive-farmer flag.",
                    "Catch-all segments (no range min) after thresholds.",
                    "New farmers never receive a segment.",
                ],
            },
            {
                "type": "heading",
                "text": "4.4 Approval Workflow",
                "level": 2,
            },
            {
                "type": "para",
                "text": "New Approval Configuration (example code: SEG_SIMULATION_ACCEPT_OFFICIAL):",
            },
            {
                "type": "bullets",
                "items": [
                    "Trigger: user submits Accept as official on simulation detail (Simulation status).",
                    "Payload: simulation reference, target season, configuration name, farmer counts, segment distribution summary.",
                    "On approve: set simulation to Official, demote other officials, replace official farmer segmentation for season.",
                    "On reject: simulation remains Simulation; notify requester.",
                    "Segregation of duties: creator and KPI editors must not be sole approver.",
                ],
            },
        ],
        "policy": [
            {
                "type": "para",
                "text": "Official segmentation affects commercial terms — dual control required. Audit simulation id, scopes, configuration, approvers, timestamps.",
            },
        ],
        "technical": [
            {
                "type": "bullets",
                "items": [
                    "Reads Farmer Contract KPI, Technologies KPI, Irregularities KPI, and Segmentation Configuration.",
                    "Writes official farmer segmentation on approval.",
                    "Integrates with Approval Workflow engine.",
                ],
            },
        ],
        "testing": [
            {
                "type": "table",
                "headers": ["ID", "Scenario", "Expected result"],
                "rows": [
                    ["T-01", "Incomplete KPI scopes", "Run blocked."],
                    ["T-02", "New farmer", "Score 0, no segment, excluded from chart."],
                    ["T-03", "Approve accept", "Official status; segmentation replaced for season."],
                    ["T-04", "Reject approval", "No segmentation change."],
                    ["T-05", "Loyalty consolidated %", "Uses kg sums not average of %."],
                ],
            },
        ],
    }
    return meta, sections


def fsd_006() -> tuple[dict, dict]:
    meta = {
        "story_id": "PA-SEG-006",
        "functionality": "Segmentation Management",
        "filename": "PA-SEG-006 - Segmentation Management.docx",
    }
    sections = {
        "process": [
            {
                "type": "bullets",
                "items": [
                    "What? Maintain official farmer segmentation per crop season; manual segment override with approval.",
                    f"Who? {OPERATOR_ROLE}; segment changes require Approval Workflow approvers.",
                    "When? After official segmentation exists for a season.",
                    "Where? Segmentation Management screen (Web).",
                    "Why? Controlled exceptions without full re-simulation.",
                    "How? Select season, review grid, propose segment change, submit for approval.",
                ],
            },
            {
                "type": "heading",
                "text": "3.2 Business rules",
                "level": 2,
            },
            {
                "type": "bullets",
                "items": [
                    "Allowed segments: only those from the configuration that produced the official snapshot.",
                    "Manual segment change does not recalculate scores (scores remain from last accepted simulation unless a new simulation is accepted).",
                    "Farmer must have official segmentation for the season to change segment.",
                    "Optional: block exclusive-only segments for non-exclusive farmers (same as simulation).",
                ],
            },
        ],
        "ui": [
            {
                "type": "heading",
                "text": "4.1 SCREEN 01: Segmentation Management (Web)",
                "level": 2,
            },
            {
                "type": "image",
                "file": "PA-SEG-006 - Manage Segmentation Wireframe.png",
                "caption": "Figure 1 — Wireframe: segmentation management by crop season.",
            },
            {
                "type": "bullets",
                "items": [
                    "Crop season selector.",
                    "Grid: farmer code, name, culture type, total score, current segment, segment dropdown, Submit for approval.",
                    "Dropdown lists only valid segments for that farmer's official configuration.",
                    "Success message after approval completes; grid refreshes.",
                ],
            },
            {
                "type": "heading",
                "text": "4.2 Approval Workflow",
                "level": 2,
            },
            {
                "type": "para",
                "text": "New Approval Configuration (example code: SEG_FARMER_SEGMENT_CHANGE):",
            },
            {
                "type": "bullets",
                "items": [
                    "Trigger: Submit for approval on a row.",
                    "Payload: farmer identity, season, current and proposed segment, total score, configuration name.",
                    "On approve: persist new segment on official farmer segmentation.",
                    "On reject: no change; notify requester.",
                    "Recommended: mandatory reason comment on request.",
                    "Approver must differ from submitter (configurable).",
                ],
            },
        ],
        "policy": [
            {
                "type": "para",
                "text": "Manual tier changes affect discounts and farmer communication — full audit. Align approver matrix with simulation accept configuration.",
            },
        ],
        "technical": [
            {
                "type": "bullets",
                "items": [
                    "Reads official farmer segmentation and configuration segments.",
                    "Approval Workflow invokes persistence on final approval.",
                ],
            },
        ],
        "testing": [
            {
                "type": "table",
                "headers": ["ID", "Scenario", "Expected result"],
                "rows": [
                    ["T-01", "Approve segment change", "New segment shown; scores unchanged."],
                    ["T-02", "Invalid segment for configuration", "Validation error at submit."],
                    ["T-03", "No official segmentation", "Cannot submit."],
                    ["T-04", "Reject approval", "Segment unchanged."],
                    ["T-05", "User without Backoffice PA User role", "No edit or submit."],
                ],
            },
        ],
    }
    return meta, sections


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    builders = [fsd_001, fsd_002, fsd_003, fsd_004, fsd_005, fsd_006]
    for build in builders:
        meta, sections = build()
        doc = _build_doc(meta, sections)
        path = OUT_DIR / meta["filename"]
        doc.save(path)
        print(f"Wrote {path}")


if __name__ == "__main__":
    main()
